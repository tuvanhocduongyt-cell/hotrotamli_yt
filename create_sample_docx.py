from docx import Document

def create_sample_docx(filename):
    doc = Document()
    doc.add_heading('MẪU ĐỀ THI ĐỊNH DẠNG CHUẨN', 0)
    
    # Câu 1: Trắc nghiệm ABCD
    doc.add_paragraph('Câu 1. Sự kiện nào đánh dấu Nguyễn Ái Quốc đã tìm thấy con đường cứu nước đúng đắn cho dân tộc Việt Nam?')
    doc.add_paragraph('A. Gửi Bản yêu sách của nhân dân An Nam tới Hội nghị Véc-xai (1919).')
    doc.add_paragraph('B. Đọc bản Sơ thảo lần thứ nhất những luận cương về vấn đề dân tộc và vấn đề thuộc địa của V.I.Lênin (1920).')
    doc.add_paragraph('C. Tham gia sáng lập Đảng Cộng sản Pháp tại Đại hội Tua (1920).')
    doc.add_paragraph('D. Thành lập Hội Việt Nam Cách mạng Thanh niên (1925).')
    doc.add_paragraph('Đáp án: B')
    
    doc.add_paragraph('') # Khoảng trống
    
    # Câu 2: Đúng/Sai có đoạn tư liệu
    doc.add_paragraph('Câu 2. Cho đoạn tư liệu sau đây:')
    doc.add_paragraph('“Ngày 3-2-1930, Hội nghị hợp nhất các tổ chức cộng sản Việt Nam đã họp tại Cửu Long (Hương Cảng, Trung Quốc) dưới sự chủ trì của Nguyễn Ái Quốc. Hội nghị đã quyết định hợp nhất các tổ chức cộng sản thành một đảng duy nhất lấy tên là Đảng Cộng sản Việt Nam.”')
    doc.add_paragraph('Dựa vào đoạn tư liệu trên, hãy chọn Đúng hoặc Sai cho các khẳng định sau:')
    doc.add_paragraph('a) Hội nghị hợp nhất các tổ chức cộng sản diễn ra vào đầu năm 1930.')
    doc.add_paragraph('b) Người chủ trì Hội nghị hợp nhất là đồng chí Trần Phú.')
    doc.add_paragraph('c) Sau hội nghị, tên gọi của Đảng là Đảng Cộng sản Đông Dương.')
    doc.add_paragraph('d) Địa điểm tổ chức hội nghị thuộc lãnh thổ Trung Quốc.')
    doc.add_paragraph('Đáp án: a-Đ, b-S, c-S, d-Đ')
    
    doc.add_paragraph('') # Khoảng trống
    
    # Câu 3: Tự luận
    doc.add_paragraph('Câu 3. Hãy phân tích ý nghĩa lịch sử của việc thành lập Đảng Cộng sản Việt Nam vào năm 1930.')
    
    doc.save(filename)
    print(f"✅ Đã tạo file: {filename}")

if __name__ == "__main__":
    create_sample_docx('mau_de_thi_chinh_xac.docx')
